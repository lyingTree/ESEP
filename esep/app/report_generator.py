# -*- coding:utf-8 -*-
"""
-------------------------------------------------------------------------------

                 Project Name : ESEP
                                                                              
                    File Name : report_generator.py

                   Start Date : 2022-03-25 07:28

                  Contributor : D.CW

                        Email : dengchuangwu@gmail.com

-------------------------------------------------------------------------------
Introduction:



-------------------------------------------------------------------------------
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt


def generate_report(case_name="Tide_cjV3.2_qiantang", save_path="./长江加密网格加了钱塘江的潮汐模型.docx"):
    station_map = "./result/{case_name}_完整区域_水深_潮位站点.png".format(case_name=case_name)
    grid_name = 'cjV3.2'
    model_info = '添加了钱塘江，径流量取15000m3/s'

    # 返回该路径下所有的 png 文件的路径
    C1_surface_path = "./result/%s/C1_surface/" % case_name
    for i in os.listdir(C1_surface_path):
        if ".png" in i:
            C1_surface = C1_surface_path + os.path.basename(i)

    C1_middle_path = "./result/%s/C1_middle/" % case_name
    for i in os.listdir(C1_middle_path):
        if ".png" in i:
            C1_middle = C1_middle_path + os.path.basename(i)

    C1_bottom_path = "./result/%s/C1_bottom/" % case_name
    for i in os.listdir(C1_bottom_path):
        if ".png" in i:
            C1_bottom = C1_bottom_path + os.path.basename(i)

    C2_surface_path = "./result/%s/C2_surface/" % case_name
    for i in os.listdir(C2_surface_path):
        if ".png" in i:
            C2_surface = C2_surface_path + os.path.basename(i)

    C2_middle_path = "./result/%s/C2_middle/" % case_name
    for i in os.listdir(C2_middle_path):
        if ".png" in i:
            C2_middle = C2_middle_path + os.path.basename(i)

    C2_bottom_path = "./result/%s/C2_bottom/" % case_name
    for i in os.listdir(C2_bottom_path):
        if ".png" in i:
            C2_bottom = C2_bottom_path + os.path.basename(i)

    C3_surface_path = "./result/%s/C3_surface/" % case_name
    for i in os.listdir(C3_surface_path):
        if ".png" in i:
            C3_surface = C3_surface_path + os.path.basename(i)
    C3_middle_path = "./result/%s/C3_middle/" % case_name
    for i in os.listdir(C3_middle_path):
        if ".png" in i:
            C3_middle = C3_middle_path + os.path.basename(i)
    C3_bottom_path = "./result/%s/C3_bottom/" % case_name
    for i in os.listdir(C3_bottom_path):
        if ".png" in i:
            C3_bottom = C3_bottom_path + os.path.basename(i)

    C4_surface_path = "./result/%s/C4_surface/" % case_name
    for i in os.listdir(C4_surface_path):
        if ".png" in i:
            C4_surface = C4_surface_path + os.path.basename(i)
    C4_middle_path = "./result/%s/C4_middle/" % case_name
    for i in os.listdir(C4_middle_path):
        if ".png" in i:
            C4_middle = C4_middle_path + os.path.basename(i)
    C4_bottom_path = "./result/%s/C4_bottom/" % case_name
    for i in os.listdir(C4_bottom_path):
        if ".png" in i:
            C4_bottom = C4_bottom_path + os.path.basename(i)

    C5_surface_path = "./result/%s/C5_surface/" % case_name
    for i in os.listdir(C5_surface_path):
        if ".png" in i:
            C5_surface = C5_surface_path + os.path.basename(i)
    C5_middle_path = "./result/%s/C5_middle/" % case_name
    for i in os.listdir(C5_middle_path):
        if ".png" in i:
            C5_middle = C5_middle_path + os.path.basename(i)
    C5_bottom_path = "./result/%s/C5_bottom/" % case_name
    for i in os.listdir(C5_bottom_path):
        if ".png" in i:
            C5_bottom = C5_bottom_path + os.path.basename(i)

    C6_surface_path = "./result/%s/C6_surface/" % case_name
    for i in os.listdir(C6_surface_path):
        if ".png" in i:
            C6_surface = C6_surface_path + os.path.basename(i)
    C6_middle_path = "./result/%s/C6_middle/" % case_name
    for i in os.listdir(C6_middle_path):
        if ".png" in i:
            C6_middle = C6_middle_path + os.path.basename(i)
    C6_bottom_path = "./result/%s/C6_bottom/" % case_name
    for i in os.listdir(C6_bottom_path):
        if ".png" in i:
            C6_bottom = C6_bottom_path + os.path.basename(i)

    C7_surface_path = "./result/%s/C7_surface/" % case_name
    for i in os.listdir(C7_surface_path):
        if ".png" in i:
            C7_surface = C7_surface_path + os.path.basename(i)
    C7_middle_path = "./result/%s/C7_middle/" % case_name
    for i in os.listdir(C7_middle_path):
        if ".png" in i:
            C7_middle = C7_middle_path + os.path.basename(i)
    C7_bottom_path = "./result/%s/C7_bottom/" % case_name
    for i in os.listdir(C7_bottom_path):
        if ".png" in i:
            C7_bottom = C7_bottom_path + os.path.basename(i)

    djs_path = "./result/%s/大戢山/" % case_name
    for i in os.listdir(djs_path):
        if ".png" in i:
            print(os.path.basename(i))
            djs = djs_path + os.path.basename(i)

    jssh_path = "./result/%s/金山石化/" % case_name
    for i in os.listdir(jssh_path):
        if ".png" in i:
            print(os.path.basename(i))
            jssh = jssh_path + os.path.basename(i)

    lcg_path = "./result/%s/芦潮港/" % case_name
    for i in os.listdir(lcg_path):
        if ".png" in i:
            print(os.path.basename(i))
            lcg = lcg_path + os.path.basename(i)

    th_path = "./result/%s/滩浒/" % case_name
    for i in os.listdir(th_path):
        if ".png" in i:
            print(os.path.basename(i))
            th = th_path + os.path.basename(i)

    zh_path = "./result/%s/镇海/" % case_name
    for i in os.listdir(zh_path):
        if ".png" in i:
            print(os.path.basename(i))
            zh = zh_path + os.path.basename(i)

    doc = Document()

    # 设置文档的基础字体
    doc.styles['Normal'].font.name = u'等线'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'等线')
    doc.styles['Normal'].font.size = Pt(14)

    # 标题
    ti = doc.add_paragraph()
    ti.paragraph_format.line_spacing = 1
    ti.paragraph_format.space_after = 1
    # 对齐方式为居中
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runt = ti.add_run("长江口模型结果")
    runt.font.size = Pt(18)
    runt.font.bold = True

    # 第一个小标题
    t1 = doc.add_paragraph()
    t1.paragraph_format.line_spacing = 1
    t1.paragraph_format.space_after = 1
    runt1 = t1.add_run("一、模型配置说明")
    runt1.font.size = Pt(16)
    runt1.font.bold = True

    # 输入网格名
    l1 = doc.add_paragraph()
    l1.paragraph_format.line_spacing = 1
    l1.paragraph_format.space_after = 1
    runl1 = l1.add_run("1.使用网格：" + grid_name)
    runl1.font.size = Pt(14)

    # 输入模型配置说明
    l2 = doc.add_paragraph()

    l2.paragraph_format.line_spacing = 1
    l2.paragraph_format.space_after = 1
    runl2 = l2.add_run("2.模型配置：" + model_info)
    runl2.font.size = Pt(14)

    doc.add_paragraph()

    # 第二个小标题
    t2 = doc.add_paragraph()
    t2.paragraph_format.line_spacing = 1
    t2.paragraph_format.space_after = 1
    runt2 = t2.add_run("二、结果分析")
    runt2.font.size = Pt(16)
    runt2.font.bold = True
    model_info_line = 1

    if len(model_info) <= 25:
        model_info_line = 1
    elif len(model_info) > 25 and len(model_info) <= 55:
        model_info_line = 2
    elif len(model_info) > 55 and len(model_info) <= 85:
        model_info_line = 3
    elif len(model_info) > 85 and len(model_info) <= 115:
        model_info_line = 4
    elif len(model_info) > 115 and len(model_info) <= 145:
        model_info_line = 5
    elif len(model_info) > 145 and len(model_info) <= 175:
        model_info_line = 6
    elif len(model_info) > 175 and len(model_info) <= 205:
        model_info_line = 7
    elif len(model_info) > 205 and len(model_info) <= 235:
        model_info_line = 8
    elif len(model_info) > 235 and len(model_info) <= 265:
        model_info_line = 9
    elif len(model_info) > 265 and len(model_info) <= 295:
        model_info_line = 10
    else:
        print('配置说明文字过多')

    # 站点图
    section = doc.sections[0]  # 获取section对象
    tl = 0.5 * 2 * 7 + 0.5 * 2 * model_info_line + 0.5 * 3 + 0.56 * 2 + 0.63

    h = section.page_height.cm - tl
    w = h * 0.933

    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(station_map, width=Cm(w), height=Cm(h))

    # 站点图说明文字
    f1 = doc.add_paragraph()
    f1.paragraph_format.line_spacing = 1
    f1.paragraph_format.space_after = 1
    # 对齐方式为居中
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runf1 = f1.add_run('Fig 1.	观测点站位图')
    runf1.font.name = '等线'
    runf1.font.size = Pt(10.5)
    runf1.font.bold = True

    # 插入分页符
    doc.add_page_break()

    # c1-c7
    # c1
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C1_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 2.	观测点C1海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C1_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 3.	观测点C1中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C1_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 4.	观测点C1底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c2
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C2_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 5.	观测点C2海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C2_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 6.	观测点C2中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C2_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = 1
    p3.paragraph_format.line_spacing = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 7.	观测点C2底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c3
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C3_surface, width=Cm(9.8), height=Cm(6.5))
    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 8.	观测点C3海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C3_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 9.	观测点C3中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C3_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 10.	观测点C3底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c4
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C4_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 11.	观测点C4海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C4_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 12.	观测点C4中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C4_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 13.	观测点C4底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c5
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C5_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 14.	观测点C5海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C5_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 15.	观测点C5中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C5_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 16.	观测点C5底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c6
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C6_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 17.	观测点C6海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C6_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 18.	观测点C6中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C6_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 19.	观测点C6底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # c7
    # 海表
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C7_surface, width=Cm(9.8), height=Cm(6.5))

    p1 = doc.add_paragraph()
    p1.paragraph_format.line_spacing = 1
    p1.paragraph_format.space_after = 1
    # 对齐方式为居中
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('Fig 20.	观测点C7海表流速、流向')
    run1.font.name = '等线'
    run1.font.size = Pt(10.5)
    # 是否加粗
    run1.font.bold = True

    # 中层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C7_middle, width=Cm(9.8), height=Cm(6.5))

    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1
    p2.paragraph_format.space_after = 1
    # 对齐方式为居中
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Fig 21.	观测点C7中层流速、流向')
    run2.font.name = '等线'
    run2.font.size = Pt(10.5)
    # 是否加粗
    run2.font.bold = True

    # 底层
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(C7_bottom, width=Cm(9.8), height=Cm(6.5))

    p3 = doc.add_paragraph()
    p3.paragraph_format.line_spacing = 1
    p3.paragraph_format.space_after = 1
    # 对齐方式为居中
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('Fig 22.	观测点C7底层流速、流向')
    run3.font.name = '等线'
    run3.font.size = Pt(10.5)
    # 是否加粗
    run3.font.bold = True

    # 大戢山
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(djs, width=Cm(9.8), height=Cm(6.5))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 1
    # 对齐方式为居中
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fig 23.	观测点大戢山潮位')
    run.font.name = '等线'
    run.font.size = Pt(10.5)
    # 是否加粗
    run.font.bold = True

    # 金山石化
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(jssh, width=Cm(9.8), height=Cm(6.5))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 1
    # 对齐方式为居中
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fig 24.	观测点金山石化潮位')
    run.font.name = '等线'
    run.font.size = Pt(10.5)
    # 是否加粗
    run.font.bold = True

    # 芦潮港
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(lcg, width=Cm(9.8), height=Cm(6.5))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 1
    # 对齐方式为居中
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fig 25.	观测点芦潮港潮位')
    run.font.name = '等线'
    run.font.size = Pt(10.5)
    # 是否加粗
    run.font.bold = True

    # 滩浒
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(th, width=Cm(9.8), height=Cm(6.5))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 1
    # 对齐方式为居中
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fig 26.	观测点滩浒潮位')
    run.font.name = '等线'
    run.font.size = Pt(10.5)
    # 是否加粗
    run.font.bold = True

    # 镇海
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_after = 1
    p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p0.add_run("")
    run.add_picture(zh, width=Cm(9.8), height=Cm(6.5))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 1
    # 对齐方式为居中
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Fig 27.	观测点镇海潮位')
    run.font.name = '等线'
    run.font.size = Pt(10.5)
    # 是否加粗
    run.font.bold = True

    # 存储文档
    doc.save(save_path)
